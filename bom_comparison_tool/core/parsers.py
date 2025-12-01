"""
Parsers for various Bill of Materials (BOM) file formats.

This module provides functions to read BOM files in different formats (XLSX, CSV,
DOCX, PDF, TXT) and normalize their content into a standardized list of BOMItem
dictionaries.

Each parser is designed to be resilient to common issues like alternate column
names and missing headers. If a file cannot be parsed, a structured error is
returned.
"""
import os
import csv
import re
from typing import List, Dict, Optional, Tuple, Any

import openpyxl
from docx import Document
from PyPDF2 import PdfReader

from .models import BOMItem, ParseResult, ErrorDict

# --- Column Name Normalization ---

# Define standard keys that the rest of the application will use.
STD_KEYS = {
    "MPN": "MPN",
    "QUANTITY": "Quantity",
    "REFDES": "RefDes",
    "DESCRIPTION": "Description",
}

# Map common variations of column names to a standard internal representation.
COLUMN_ALIASES = {
    "MPN": ["mpn", "part number", "manufacturer part number", "mfg part number"],
    "QUANTITY": ["quantity", "qty", "quant"],
    "REFDES": ["refdes", "reference designator", "designator", "ref des"],
    "DESCRIPTION": ["description", "desc"],
}

def _normalize_header(header: List[str]) -> Dict[str, str]:
    """Converts a raw header list to a standardized dictionary."""
    normalized = {}
    for column in header:
        column_clean = str(column).lower().strip()
        for std_key, aliases in COLUMN_ALIASES.items():
            if column_clean in aliases:
                normalized[column] = STD_KEYS[std_key]
                break
    return normalized

def _find_header_map(rows: List[List[Any]]) -> Optional[Tuple[int, Dict[int, str]]]:
    """
    Identifies the header row and creates a map from column index to standard key.

    Args:
        rows: A list of rows, where each row is a list of cell values.

    Returns:
        A tuple containing (header_row_index, {column_index: standard_key}),
        or None if a valid header cannot be found.
    """
    for i, row in enumerate(rows):
        # Filter out potential None values from empty cells
        row_values = [str(cell).strip() for cell in row if cell is not None]
        
        normalized_headers = _normalize_header(row_values)
        
        # Consider it a valid header if at least two standard keys are found.
        # This is a heuristic to avoid false positives.
        if len(normalized_headers) >= 2:
            header_map = {}
            for j, cell in enumerate(row):
                std_key = normalized_headers.get(str(cell).strip())
                if std_key:
                    header_map[j] = std_key
            return i, header_map
    return None

def _process_data_rows(rows: List[List[Any]], header_map: Dict[int, str], start_index: int) -> List[BOMItem]:
    """Processes rows into a list of BOMItem dictionaries using the header map."""
    bom_items = []
    for row in rows[start_index:]:
        if not any(row):  # Skip empty rows
            continue

        item_data = {}
        for idx, key in header_map.items():
            if idx < len(row):
                item_data[key] = str(row[idx]).strip() if row[idx] is not None else ""

        # --- Data Cleaning and Type Conversion ---
        # Ensure all required keys are present, defaulting to empty values
        mpn = item_data.get("MPN", "")
        desc = item_data.get("Description", "")
        
        # Skip rows that don't have a Manufacturer Part Number
        if not mpn:
            continue

        try:
            quantity = int(item_data.get("Quantity", 0))
        except (ValueError, TypeError):
            quantity = 0 # Default to 0 if conversion fails

        refdes_str = item_data.get("RefDes", "")
        # Split RefDes by common delimiters (comma, space, semicolon)
        refdes = [r.strip() for r in re.split(r'[,;\s]+', refdes_str) if r.strip()]

        bom_items.append(BOMItem(
            MPN=mpn,
            Quantity=quantity,
            RefDes=refdes,
            Description=desc
        ))
    return bom_items

# --- Individual File Parsers ---

def _parse_xlsx(file_path: str) -> ParseResult:
    """Parses an XLSX file."""
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        sheet = workbook.active
        rows = [list(cell.value for cell in row) for row in sheet.iter_rows()]
        
        header_info = _find_header_map(rows)
        if not header_info:
            return ErrorDict(error="Could not find a valid header row.", file_path=file_path)
        
        header_idx, header_map = header_info
        return _process_data_rows(rows, header_map, header_idx + 1)

    except Exception as e:
        return ErrorDict(error=f"Failed to parse XLSX file: {e}", file_path=file_path)

def _parse_csv(file_path: str) -> ParseResult:
    """Parses a CSV file."""
    try:
        with open(file_path, mode='r', encoding='utf-8-sig') as infile:
            # Sniff to detect the dialect (comma, tab, etc.)
            dialect = csv.Sniffer().sniff(infile.read(1024))
            infile.seek(0)
            reader = csv.reader(infile, dialect)
            rows = list(reader)

        header_info = _find_header_map(rows)
        if not header_info:
            return ErrorDict(error="Could not find a valid header row.", file_path=file_path)
            
        header_idx, header_map = header_info
        return _process_data_rows(rows, header_map, header_idx + 1)

    except Exception as e:
        return ErrorDict(error=f"Failed to parse CSV file: {e}", file_path=file_path)

def _parse_txt(file_path: str) -> ParseResult:
    """
    Parses a TXT file, treating it like a CSV with an unknown delimiter.
    This is functionally similar to the CSV parser but for .txt extensions.
    """
    try:
        with open(file_path, mode='r', encoding='utf-8') as infile:
            # Read a sample to sniff for delimiter
            sample = infile.read(2048)
            infile.seek(0)
            
            # Simple sniffer: check for tab, then comma, then semicolon
            if '\t' in sample:
                delimiter = '\t'
            elif ',' in sample:
                delimiter = ','
            elif ';' in sample:
                delimiter = ';'
            else: # Fallback to splitting by multiple spaces
                delimiter = None

            if delimiter:
                reader = csv.reader(infile, delimiter=delimiter)
                rows = list(reader)
            else:
                lines = infile.readlines()
                rows = [re.split(r'\s{2,}', line.strip()) for line in lines]

        header_info = _find_header_map(rows)
        if not header_info:
            return ErrorDict(error="Could not find a valid header row.", file_path=file_path)
            
        header_idx, header_map = header_info
        return _process_data_rows(rows, header_map, header_idx + 1)

    except Exception as e:
        return ErrorDict(error=f"Failed to parse TXT file: {e}", file_path=file_path)

def _parse_docx(file_path: str) -> ParseResult:
    """Parses tables from a DOCX file."""
    try:
        document = Document(file_path)
        all_items = []
        if not document.tables:
            return ErrorDict(error="No tables found in the DOCX file.", file_path=file_path)

        for table in document.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            header_info = _find_header_map(rows)
            if header_info:
                header_idx, header_map = header_info
                all_items.extend(_process_data_rows(rows, header_map, header_idx + 1))
        
        if not all_items:
            return ErrorDict(error="Found tables but could not extract valid BOM data.", file_path=file_path)
            
        return all_items

    except Exception as e:
        return ErrorDict(error=f"Failed to parse DOCX file: {e}", file_path=file_path)

def _parse_pdf(file_path: str) -> ParseResult:
    """
    Parses a text-based PDF file.
    
    Note: This is a best-effort parser. PyPDF2 extracts raw text, so this
    function attempts to reconstruct rows and columns. It is not as reliable
    as table-aware libraries like pdfplumber and will fail on complex layouts.
    """
    try:
        reader = PdfReader(file_path)
        text_content = ""
        for page in reader.pages:
            text_content += page.extract_text() + "\n"

        # Treat the extracted text as a single block and parse like a TXT file
        lines = text_content.strip().split('\n')
        # Assume columns are separated by two or more spaces
        rows = [re.split(r'\s{2,}', line.strip()) for line in lines if line.strip()]

        header_info = _find_header_map(rows)
        if not header_info:
            return ErrorDict(error="Could not find a valid header in the extracted PDF text.", file_path=file_path)
        
        header_idx, header_map = header_info
        return _process_data_rows(rows, header_map, header_idx + 1)

    except Exception as e:
        return ErrorDict(error=f"Failed to parse PDF file: {e}", file_path=file_path)


# --- Main Dispatcher Function ---

def parse_bom_file(file_path: str) -> ParseResult:
    """
    Parses a BOM file, dispatching to the correct parser based on extension.

    Args:
        file_path: The absolute or relative path to the BOM file.

    Returns:
        A ParseResult, which is either a list of BOMItem dictionaries on success
        or an ErrorDict on failure.
    """
    if not os.path.exists(file_path):
        return ErrorDict(error="File not found.", file_path=file_path)

    _, extension = os.path.splitext(file_path.lower())

    if extension == '.xlsx':
        return _parse_xlsx(file_path)
    if extension == '.csv':
        return _parse_csv(file_path)
    if extension == '.txt':
        return _parse_txt(file_path)
    if extension == '.docx':
        return _parse_docx(file_path)
    if extension == '.pdf':
        return _parse_pdf(file_path)

    return ErrorDict(error=f"Unsupported file extension: '{extension}'", file_path=file_path)
